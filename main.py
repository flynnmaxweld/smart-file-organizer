import sys
import os
import shutil
import json
import time
import hashlib
import re  # Added for Regex matching
import zipfile  # Added for backup feature
from PIL import Image, UnidentifiedImageError
# Add PIL.ExifTags for accessing metadata by name
from PIL.ExifTags import TAGS
from datetime import datetime  # removed unused timedelta

# --- Qt Imports ---
from PySide6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QPushButton, QLineEdit, QTextEdit, QProgressBar, QCheckBox,
    QFileDialog, QMessageBox, QLabel,
    QTableWidget, QTableWidgetItem, QHeaderView, QDialog, QComboBox,
    QSystemTrayIcon, QMenu, QFrame
)
from PySide6.QtCore import QObject, Signal, QThread, Slot, Qt, QPoint
from PySide6.QtGui import QColor, QIcon, QAction, QDropEvent, QDragEnterEvent, QFont

# --- Icon Library ---
import qtawesome as qta

# --- Optional Imports for Advanced Features ---
try:
    from watchdog.observers import Observer
    from watchdog.events import FileSystemEventHandler
    WATCHDOG_AVAILABLE = True
except ImportError:
    WATCHDOG_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


# --- Constants ---
LOG_FILE_NAME = ".organizer_log.json"  # Still used for the single 'undo' log
SETTINGS_FILE_NAME = ".organizer_settings.json"  # Still used for theme, rules, categories

# --- Aesthetic Stylesheets (QSS) ---
DARK_BLUE_THEME = """
    #MainWindow {
        background-color: #0d1117; /* GitHub dark background */
        border-radius: 10px;
    }
    #CentralWidget { background-color: transparent; }
    #TitleBar { background-color: transparent; }
    #WindowButton { background-color: transparent; border: none; padding: 5px; border-radius: 6px; }
    #WindowButton:hover { background-color: #30363d; } /* GitHub dark border/hover */
    #CloseButton:hover { background-color: #e81123; }

    /* --- Main Content Card (Now the main container) --- */
    #ContentCard {
        background-color: #161b22; /* GitHub dark lighter background */
        border-radius: 8px; /* Consistent radius */
    }

    /* --- Toolbar Styles --- */
    #ToolBar {
        background-color: #0d1117; /* Match main background */
        border-top-left-radius: 8px; /* Match card radius */
        border-top-right-radius: 8px;
        padding: 8px 10px; /* Adjusted padding */
        border-bottom: 1px solid #30363d;
    }
    #ToolBar QLabel { color: #8b949e; /* GitHub dark secondary text */ font-weight: bold; background: transparent; font-size: 8pt; } /* Smaller label */
    #ToolBar QLineEdit {
        background-color: #0d1117; border: 1px solid #30363d; border-radius: 6px;
        padding: 6px 9px; /* Adjusted padding */
        color: #c9d1d9; font-size: 9pt;
    }
    #ToolBar QPushButton {
        background-color: #21262d; color: #c9d1d9; border: 1px solid #30363d;
        padding: 6px 10px; /* Adjusted padding */
        border-radius: 6px; font-size: 9pt; font-weight: 500; /* Medium weight */
    }
    #ToolBar QPushButton:hover { background-color: #30363d; border-color: #8b949e; }
    #ToolBar QPushButton#OrganizeBtn { /* Specific style for Organize button */
        color: white; background-color: #238636; border-color: #3fb950; font-weight: bold; /* GitHub green */
    }
     #ToolBar QPushButton#OrganizeBtn:hover { background-color: #2ea043; }
    #ToolBar QFrame { background-color: #30363d; max-width: 1px;} /* Thinner separator */
    #ToolBar QComboBox {
        background-color: #21262d; color: #c9d1d9; border: 1px solid #30363d;
        padding: 5px 8px; /* Adjusted padding */
        border-radius: 6px; font-size: 9pt; min-width: 60px;
    }
    #ToolBar QComboBox::drop-down { border: none; width: 15px; }
    #ToolBar QComboBox QAbstractItemView { background-color: #161b22; border: 1px solid #30363d; selection-background-color: #007AFF; } /* Keep blue selection */
    #ThemeToggleBtn, #HelpBtn { /* Small icon buttons */
        padding: 5px; border: none; background-color: transparent; border-radius: 6px;
    }
    #ThemeToggleBtn:hover, #HelpBtn:hover { background-color: #30363d; }
    /* Style for Watcher/Backup Toggles */
    #ToolBar QCheckBox { color: #8b949e; font-size: 8pt; margin-left: 5px;}
    #ToolBar QCheckBox::indicator { width: 14px; height: 14px; background-color: #0d1117; border: 1px solid #30363d; border-radius: 3px;}
    #ToolBar QCheckBox::indicator:checked { background-color: #238636; border: 1px solid #3fb950;}
    #ToolBar QCheckBox#WatcherToggle::indicator { width: 34px; height: 18px; background-color: #30363d; border-radius: 9px; border: 1px solid #444; }
    #ToolBar QCheckBox#WatcherToggle::indicator:checked { background-color: #238636; } /* GitHub green */
    #ToolBar QCheckBox#WatcherToggle::indicator::handle { background-color: #8b949e; width: 14px; height: 14px; border-radius: 7px; }
    #ToolBar QCheckBox#WatcherToggle::indicator::handle:unchecked { subcontrol-position: left; margin: 1px; }
    #ToolBar QCheckBox#WatcherToggle::indicator::handle:checked { subcontrol-position: right; margin: 1px; }


    /* --- Content Area --- */
    #ContentArea { background-color: transparent; }
    QWidget { color: #c9d1d9; font-family: 'Segoe UI', sans-serif; font-size: 9pt; } /* Base font size */
    #StatsLabel { color: #8b949e; font-size: 8pt; padding: 0 10px 5px 10px; background-color: #0d1117; } /* Stats styling */
    #LogToolsWidget QPushButton { /* Style for Export Log button */
        background-color: #21262d; color: #c9d1d9; border: 1px solid #30363d;
        padding: 4px 8px; border-radius: 6px; font-size: 8pt;
    }
    #LogToolsWidget QPushButton:hover { background-color: #30363d; }
    
    QTextEdit {
        background-color: #0d1117; /* Match main background */
        border: none;
        color: #8b949e; /* GitHub dark secondary text for logs */
        font-size: 8.5pt; /* Slightly smaller log text */
        padding: 10px;
        border-radius: 6px;
    }
    QProgressBar { border: none; text-align: center; height: 5px; background-color: #30363d; border-radius: 2.5px; margin: 5px 10px 8px 10px; } /* Thinner bar */
    QProgressBar::chunk { background-color: #238636; border-radius: 2.5px; } /* GitHub green */

    /* --- Dialog Styles --- */
    QDialog { background-color: #161b22; border: 1px solid #30363d; }
    QDialog QLineEdit, QDialog QComboBox, QDialog QTableWidget, QDialog QSpinBox { /* Added SpinBox */
        background-color: #0d1117; border: 1px solid #30363d;
        border-radius: 6px; padding: 8px; color: #c9d1d9; font-size: 9pt;
    }
    QDialog QHeaderView::section { background-color: #21262d; padding: 5px; border: none; color: #c9d1d9; font-weight: bold;}
    QDialog QPushButton { padding: 8px 16px; background-color: #21262d; border-radius: 6px; border: 1px solid #30363d; color: #c9d1d9;}
    QDialog QPushButton:hover { background-color: #30363d; }
    QDialog QPushButton#SaveButton, QDialog QPushButton#ApproveButton { background-color: #238636; color: #ffffff; border: 1px solid #3fb950;} /* GitHub green */
    QDialog QPushButton#DeleteButton { background-color: #da3633; border: 1px solid #ff7b72; color: white; } /* GitHub red */
    QLabel#DeleteWarningLabel { color: #f85149; font-weight: bold; } /* Warning label style */
"""

LIGHT_BLUE_THEME = """
    #MainWindow {
        background-color: #f6f8fa; /* GitHub light background */
        border-radius: 10px;
        border: 1px solid #d0d7de;
    }
    #CentralWidget { background-color: transparent; }
    #TitleBar { background-color: transparent; }
    #WindowButton { background-color: transparent; border: none; padding: 5px; border-radius: 6px; }
    #WindowButton:hover { background-color: #e1e4e8; }
    #CloseButton:hover { background-color: #ef4444; }

    /* --- Main Content Card --- */
    #ContentCard { background-color: #ffffff; border-radius: 8px; }

    /* --- Toolbar Styles --- */
    #ToolBar {
        background-color: #f6f8fa; border-top-left-radius: 8px; border-top-right-radius: 8px;
        padding: 8px 10px; border-bottom: 1px solid #d0d7de; /* GitHub light border */
    }
     #ToolBar QLabel { color: #57606a; /* GitHub light secondary text */ font-weight: bold; background: transparent; font-size: 8pt; }
    #ToolBar QLineEdit {
        background-color: #ffffff; border: 1px solid #d0d7de; border-radius: 6px;
        padding: 6px 9px; color: #24292f; /* GitHub light primary text */ font-size: 9pt;
    }
    #ToolBar QPushButton {
        background-color: #f6f8fa; color: #24292f; border: 1px solid #d0d7de;
        padding: 6px 10px; border-radius: 6px; font-size: 9pt; font-weight: 500;
    }
    #ToolBar QPushButton:hover { background-color: #e1e4e8; border-color: #afb8c1; }
    #ToolBar QPushButton#OrganizeBtn {
        color: white; background-color: #1f883d; border-color: #269f47; font-weight: bold; /* GitHub light green */
    }
     #ToolBar QPushButton#OrganizeBtn:hover { background-color: #269f47; }
     #ToolBar QFrame { background-color: #d0d7de; max-width: 1px;}
     #ToolBar QComboBox {
        background-color: #f6f8fa; color: #24292f; border: 1px solid #d0d7de;
        padding: 5px 8px; border-radius: 6px; font-size: 9pt; min-width: 60px;
     }
     #ToolBar QComboBox::drop-down { border: none; width: 15px; }
     #ToolBar QComboBox QAbstractItemView { background-color: #ffffff; border: 1px solid #d0d7de; selection-background-color: #0969da; } /* GitHub light blue selection */
     #ThemeToggleBtn, #HelpBtn { padding: 5px; border: none; background-color: transparent; border-radius: 6px;}
     #ThemeToggleBtn:hover, #HelpBtn:hover { background-color: #e1e4e8; }
     /* Style for Watcher/Backup Toggles */
     #ToolBar QCheckBox { color: #57606a; font-size: 8pt; margin-left: 5px;}
     #ToolBar QCheckBox::indicator { width: 14px; height: 14px; background-color: #f6f8fa; border: 1px solid #d0d7de; border-radius: 3px;}
     #ToolBar QCheckBox::indicator:checked { background-color: #1f883d; border: 1px solid #269f47;}
     #ToolBar QCheckBox#WatcherToggle::indicator { width: 34px; height: 18px; background-color: #d0d7de; border-radius: 9px; border: 1px solid #afb8c1; }
     #ToolBar QCheckBox#WatcherToggle::indicator:checked { background-color: #1f883d; } /* GitHub light green */
     #ToolBar QCheckBox#WatcherToggle::indicator::handle { background-color: #ffffff; width: 14px; height: 14px; border-radius: 7px; }
     #ToolBar QCheckBox#WatcherToggle::indicator::handle:unchecked { subcontrol-position: left; margin: 1px; }
     #ToolBar QCheckBox#WatcherToggle::indicator::handle:checked { subcontrol-position: right; margin: 1px; }

    /* --- Content Area --- */
    #ContentArea { background-color: transparent; }
    QWidget { color: #24292f; font-family: 'Segoe UI', sans-serif; font-size: 9pt; }
    #StatsLabel { color: #57606a; font-size: 8pt; padding: 0 10px 5px 10px; background-color: #f6f8fa; } /* Stats styling */
    #LogToolsWidget QPushButton { /* Style for Export Log button */
        background-color: #f6f8fa; color: #24292f; border: 1px solid #d0d7de;
        padding: 4px 8px; border-radius: 6px; font-size: 8pt;
    }
    #LogToolsWidget QPushButton:hover { background-color: #e1e4e8; }

    QTextEdit { background-color: #f6f8fa; border: none; color: #57606a; font-size: 8.5pt; padding: 10px; border-radius: 6px;}
    QProgressBar { border: none; text-align: center; height: 5px; background-color: #e1e4e8; border-radius: 2.5px; margin: 5px 10px 8px 10px; }
    QProgressBar::chunk { background-color: #1f883d; border-radius: 2.5px; } /* GitHub light green */

    /* --- Dialog Styles --- */
    QDialog { background-color: #ffffff; border: 1px solid #d0d7de; }
    QDialog QLineEdit, QDialog QComboBox, QDialog QTableWidget, QDialog QSpinBox { /* Added SpinBox */
        background-color: #f6f8fa; border: 1px solid #d0d7de; border-radius: 6px; padding: 8px; color: #24292f; font-size: 9pt;
    }
    QDialog QHeaderView::section { background-color: #f6f8fa; padding: 5px; border: none; color: #57606a; font-weight: bold;}
    QDialog QPushButton { padding: 8px 16px; background-color: #f6f8fa; border-radius: 6px; border: 1px solid #d0d7de; color: #24292f;}
    QDialog QPushButton:hover { background-color: #e1e4e8; }
    QDialog QPushButton#SaveButton, QDialog QPushButton#ApproveButton { background-color: #1f883d; color: white; border: 1px solid #269f47;} /* GitHub light green */
    QDialog QPushButton#DeleteButton { background-color: #cf222e; border: 1px solid #ff7b72; color: white; } /* GitHub light red */
    QLabel#DeleteWarningLabel { color: #cf222e; font-weight: bold; } /* Warning label style */
"""


class Worker(QObject):
    log_message = Signal(str)
    progress_updated = Signal(int, int)
    finished = Signal()
    single_file_organized = Signal(str)
    duplicate_scan_finished = Signal(dict)
    organization_preview_ready = Signal(list, dict, bool)
    folder_stats_ready = Signal(int, str, dict)  # MODIFIED: int, str, dict

    def __init__(self):
        super().__init__()
        self.categories = {}
        self.rules = []
        self.conflict_strategy = "rename"  # Default strategy
        self.file_metadata_cache = {}  # Cache for EXIF/Doc metadata
        self.has_content_rules = False  # ADDED: Flag for content scan optimization
        self.load_settings()  # Load initial settings

    def load_settings(self):
        """Loads settings from the default JSON file."""
        try:
            with open(SETTINGS_FILE_NAME, "r") as f:
                settings = json.load(f)
                self.categories = settings.get("categories", self._get_defaults("categories"))
                self.rules = settings.get("rules", self._get_defaults("rules"))
                # ADDED: Check for content rules
                self.has_content_rules = any(
                    cond.get("type") == "content" for rule in self.rules for cond in rule.get("conditions", [])
                )
                loaded_conflict = settings.get("conflict_strategy", "rename")
                self.conflict_strategy = "rename" if loaded_conflict == "overwrite" else loaded_conflict  # Migrate away from overwrite
        except (FileNotFoundError, json.JSONDecodeError):
            self.categories = self._get_defaults("categories")
            self.rules = self._get_defaults("rules")
            # ADDED: Check for content rules
            self.has_content_rules = any(
                cond.get("type") == "content" for rule in self.rules for cond in rule.get("conditions", [])
            )
            self.conflict_strategy = "rename"  # Ensure default on error

    def _get_defaults(self, key):
        if key == "categories":
            # ... (Default categories remain the same)
            return {
                "Images": [
                    ".jpg",
                    ".jpeg",
                    ".png",
                    ".gif",
                    ".bmp",
                    ".tiff",
                    ".tif",
                    ".svg",
                    ".webp",
                    ".heic",
                    ".heif",
                    ".raw",
                    ".cr2",
                    ".nef",
                    ".arw",
                    ".orf",
                    ".sr2",
                    ".psd",
                    ".ai",
                    ".eps",
                    ".indd",
                ],
                "Videos": [".mp4", ".mkv", ".flv", ".avi", ".mov", ".wmv", ".webm", ".mpg", ".mpeg", ".3gp", ".m4v"],
                "Audio": [".mp3", ".wav", ".aac", ".flac", ".ogg", ".m4a", ".wma", ".aiff"],
                "Documents": [".pdf", ".doc", ".docx", ".txt", ".rtf", ".odt", ".pages", ".tex", ".md"],
                "Spreadsheets": [".xls", ".xlsx", ".csv", ".ods", ".numbers"],
                "Presentations": [".ppt", ".pptx", ".odp", ".key"],
                "Archives": [".zip", ".rar", ".7z", ".tar", ".gz", ".bz2", ".xz", ".iso", ".dmg"],
                "Executables": [".exe", ".msi", ".bat", ".sh", ".app", ".jar", ".com"],
                "Code": [
                    ".py",
                    ".js",
                    ".ts",
                    ".jsx",
                    ".tsx",
                    ".html",
                    ".css",
                    ".scss",
                    ".java",
                    ".cpp",
                    ".c",
                    ".h",
                    ".cs",
                    ".php",
                    ".rb",
                    ".go",
                    ".swift",
                    ".kt",
                    ".json",
                    ".xml",
                    ".yaml",
                    ".yml",
                    ".sql",
                ],
                "Web": [".html", ".htm", ".css", ".js", ".php", ".asp", ".aspx"],
                "Fonts": [".ttf", ".otf", ".woff", ".woff2", ".eot"],
                "Data": [".csv", ".json", ".xml", ".yaml", ".yml", ".sql", ".db", ".sqlite", ".mdb"],
                "System": [".dll", ".sys", ".drv", ".ini", ".cfg"],
                "Other": [],
            }
        if key == "rules":
            # ... (Default rules remain the same)
            return [
                {
                    "name": "Sort Invoices into Finances",
                    "conditions": [{"type": "content", "value": "invoice"}],
                    "action_type": "move",
                    "action_value": "Finances/{YYYY-MM}",
                },
                {
                    "name": "Sort Reports",
                    "conditions": [{"type": "filename", "value": "report"}],
                    "action_type": "move",
                    "action_value": "Documents/Reports",
                },
            ]
        return {}

    # --- Metadata Helper Functions ---
    # ... (_get_exif_data, _get_pdf_metadata, _get_docx_metadata remain the same) ...
    def _get_exif_data(self, file_path):
        """Extracts and decodes EXIF data from an image file."""
        if not file_path.lower().endswith(tuple(self.categories.get("Images", []))):
            return {}
        try:
            with Image.open(file_path) as img:
                exif_data = img._getexif()
                if not exif_data:
                    return {}

                exif = {}
                for key, val in exif_data.items():
                    tag = TAGS.get(key, key)
                    if isinstance(val, bytes):
                        try:
                            val = val.decode("utf-8", errors="ignore").strip()
                        except Exception:
                            val = str(val)
                    if isinstance(val, str):
                        val = val.strip()
                    exif[tag] = val

                # Special handling for FNumber
                if "FNumber" in exif:
                    f_num = exif["FNumber"]
                    if isinstance(f_num, tuple) and len(f_num) == 2 and f_num[1] != 0:
                        exif["FNumber"] = float(f_num[0] / f_num[1])
                    elif isinstance(f_num, (int, float)):
                        exif["FNumber"] = float(f_num)

                return exif
        except (UnidentifiedImageError, IOError, AttributeError, OSError, TypeError):
            # Catch more errors as EXIF data can be corrupt
            return {}

    def _get_pdf_metadata(self, file_path):
        """Extracts metadata from a PDF file."""
        if not PYPDF2_AVAILABLE or not file_path.lower().endswith(".pdf"):
            return {}
        try:
            with open(file_path, "rb") as f:
                reader = PyPDF2.PdfReader(f, strict=False)
                if reader.is_encrypted:
                    return {}
                info = reader.metadata
                return {"author": getattr(info, "author", "") or "", "title": getattr(info, "title", "") or ""}
        except Exception:
            return {}

    def _get_docx_metadata(self, file_path):
        """Extracts metadata from a DOCX file."""
        if not DOCX_AVAILABLE or not file_path.lower().endswith(".docx"):
            return {}
        try:
            doc = docx.Document(file_path)
            props = doc.core_properties
            return {"author": props.author or "", "title": props.title or ""}
        except Exception:
            return {}

    # --- Core Logic ---
    # ... (check_file_content, check_rule_conditions, determine_destination_and_action remain the same) ...
    def check_file_content(self, file_path, keyword):
        keyword = keyword.lower()
        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".txt":
                with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                    content = f.read(1024 * 1024)
                if keyword in content.lower():
                    return True
            elif PYPDF2_AVAILABLE and ext == ".pdf":
                if os.path.getsize(file_path) > 5 * 1024 * 1024:
                    return False
                with open(file_path, "rb") as f:
                    reader = PyPDF2.PdfReader(f, strict=False)
                if reader.is_encrypted:
                    return False
                for i in range(min(len(reader.pages), 2)):
                    page = reader.pages[i]
                    page_text = page.extract_text()
                    if page_text and keyword in page_text.lower():
                        return True
            elif DOCX_AVAILABLE and ext == ".docx":
                doc = docx.Document(file_path)
                for i in range(min(len(doc.paragraphs), 10)):
                    para = doc.paragraphs[i]
                    if keyword in para.text.lower():
                        return True
        except Exception as exc:
            self.log_message.emit(f"  - WARN: Scan content error {os.path.basename(file_path)}: {exc}")
            return False
        return False

    def check_rule_conditions(self, file_path, conditions):
        """Checks if a file meets all conditions of a rule."""
        if not conditions:
            return False

        try:
            file_size_mb = os.path.getsize(file_path) / (1024 * 1024)
            file_mtime = datetime.fromtimestamp(os.path.getmtime(file_path))
            now = datetime.now()
        except OSError as exc:
            self.log_message.emit(f"  - WARN: Could not get stats for {os.path.basename(file_path)}: {exc}")
            return False

        # Ensure file cache exists for this path
        if file_path not in self.file_metadata_cache:
            self.file_metadata_cache[file_path] = {}

        for cond in conditions:
            cond_type = cond.get("type", "").lower()
            cond_value = cond.get("value", "")
            cond_matcher = cond.get("matcher", "contains").lower()
            if not cond_type or not cond_value:
                continue

            match = False

            try:
                if cond_type == "filename":
                    filename = os.path.basename(file_path)
                    if cond_matcher == "contains" and cond_value.lower() in filename.lower():
                        match = True
                    elif cond_matcher == "startswith" and filename.lower().startswith(cond_value.lower()):
                        match = True
                    elif cond_matcher == "endswith" and filename.lower().endswith(cond_value.lower()):
                        match = True
                    elif cond_matcher == "equals" and filename.lower() == cond_value.lower():
                        match = True
                    elif cond_matcher == "regex" and re.search(cond_value, filename, re.IGNORECASE):
                        match = True

                elif cond_type == "content":
                    # --- OPTIMIZATION ---
                    if not self.has_content_rules:
                        match = False  # Skip if no content rules exist globally
                    elif self.check_file_content(file_path, cond_value):
                        match = True
                    # --- END OPTIMIZATION ---

                elif cond_type == "filesize_mb":
                    size_val = float(cond_value)
                    if cond_matcher == "greater_than" and file_size_mb > size_val:
                        match = True
                    elif cond_matcher == "less_than" and file_size_mb < size_val:
                        match = True

                elif cond_type == "date_modified_days":
                    days_val = int(cond_value)
                    delta = now - file_mtime
                    if cond_matcher == "older_than" and delta.days > days_val:
                        match = True
                    elif cond_matcher == "newer_than" and delta.days < days_val:
                        match = True

                elif cond_type == "original folder name":
                    folder_name = os.path.basename(os.path.dirname(file_path))
                    if cond_matcher == "contains" and cond_value.lower() in folder_name.lower():
                        match = True
                    elif cond_matcher == "startswith" and folder_name.lower().startswith(cond_value.lower()):
                        match = True
                    elif cond_matcher == "endswith" and folder_name.lower().endswith(cond_value.lower()):
                        match = True
                    elif cond_matcher == "equals" and folder_name.lower() == cond_value.lower():
                        match = True
                    elif cond_matcher == "regex" and re.search(cond_value, folder_name, re.IGNORECASE):
                        match = True

                elif cond_type in ["camera model", "lens model", "artist"]:
                    if "exif" not in self.file_metadata_cache[file_path]:
                        self.file_metadata_cache[file_path]["exif"] = self._get_exif_data(file_path)
                    exif = self.file_metadata_cache[file_path]["exif"]
                    if not exif:
                        continue

                    key = {"camera model": "Model", "lens model": "LensModel", "artist": "Artist"}.get(cond_type)
                    value = str(exif.get(key, ""))

                    if cond_matcher == "contains" and cond_value.lower() in value.lower():
                        match = True
                    elif cond_matcher == "equals" and cond_value.lower() == value.lower():
                        match = True

                elif cond_type == "f-stop":
                    if "exif" not in self.file_metadata_cache[file_path]:
                        self.file_metadata_cache[file_path]["exif"] = self._get_exif_data(file_path)
                    exif = self.file_metadata_cache[file_path]["exif"]
                    if not exif or "FNumber" not in exif:
                        continue

                    value = exif.get("FNumber")
                    cond_num = float(cond_value)

                    if cond_matcher == "equals" and value == cond_num:
                        match = True
                    elif cond_matcher == "greater_than" and value > cond_num:
                        match = True
                    elif cond_matcher == "less_than" and value < cond_num:
                        match = True

                elif cond_type in ["pdf author", "pdf title", "docx author", "docx title"]:
                    if "doc_meta" not in self.file_metadata_cache[file_path]:
                        if file_path.lower().endswith(".pdf"):
                            self.file_metadata_cache[file_path]["doc_meta"] = self._get_pdf_metadata(file_path)
                        elif file_path.lower().endswith(".docx"):
                            self.file_metadata_cache[file_path]["doc_meta"] = self._get_docx_metadata(file_path)
                        else:
                            self.file_metadata_cache[file_path]["doc_meta"] = {}

                    meta = self.file_metadata_cache[file_path].get("doc_meta", {})
                    if not meta:
                        continue

                    key = "author" if "author" in cond_type else "title"
                    value = str(meta.get(key, ""))

                    if cond_matcher == "contains" and cond_value.lower() in value.lower():
                        match = True
                    elif cond_matcher == "equals" and cond_value.lower() == value.lower():
                        match = True

            except Exception as exc:
                # Catch errors from regex, float conversion, etc.
                self.log_message.emit(
                    f"  - WARN: Rule check error for {os.path.basename(file_path)} (Type: {cond_type}, Value: {cond_value}): {exc}"
                )
                continue  # Skip this condition if it fails

            if not match:
                return False  # If any condition fails, exit

        return True  # All conditions passed

    def determine_destination_and_action(self, file_path):
        """Determines destination folder OR new name based on rules/categories."""
        # Ensure cache exists for this file
        if file_path not in self.file_metadata_cache:
            self.file_metadata_cache[file_path] = {}

        # 1. Check Rules first
        for rule in self.rules:
            if self.check_rule_conditions(file_path, rule.get("conditions")):
                action_type = rule.get("action_type", "move").lower()
                action_value = rule.get("action_value", "")

                # 'delete' action doesn't require a value
                if action_type == "delete":
                    return "delete", None, None  # No destination/name needed

                # Other actions require a value
                if not action_value:
                    continue

                now = datetime.now()
                formatted_value = (
                    action_value.replace("{YYYY}", now.strftime("%Y")).replace("{MM}", now.strftime("%m")).replace("{DD}", now.strftime("%d"))
                )

                if action_type == "rename":
                    orig_name, orig_ext = os.path.splitext(os.path.basename(file_path))
                    new_name = formatted_value.replace("{OrigName}", orig_name).replace("{Ext}", orig_ext).replace("{SeqNum}", "1")
                    if not new_name.endswith(orig_ext):
                        new_name += orig_ext
                    return "rename", os.path.dirname(file_path), new_name
                elif action_type == "move" or action_type == "copy":  # Handle move and copy similarly here
                    return action_type, formatted_value, os.path.basename(file_path)

        # 2. Check Categories (default sorting, always 'move')
        dest_folder = "Other"
        ext = os.path.splitext(file_path)[1].lower()
        if ext in self.categories.get("Images", []):
            try:
                # Use cached EXIF data if available
                if "exif" not in self.file_metadata_cache[file_path]:
                    self.file_metadata_cache[file_path]["exif"] = self._get_exif_data(file_path)
                exif = self.file_metadata_cache[file_path].get("exif", {})

                if exif:
                    dt_str = exif.get("DateTimeOriginal") or exif.get("DateTime")
                    if dt_str:
                        try:
                            # dt_str is already decoded string
                            date_obj = datetime.strptime(str(dt_str), "%Y:%m:%d %H:%M:%S")
                            dest_folder = os.path.join("Images", date_obj.strftime("%Y-%m"))
                        except (ValueError, TypeError):
                            pass

                    desc = str(exif.get("ImageDescription", ""))
                    if "screenshot" in desc.lower():
                        dest_folder = os.path.join("Images", "Screenshots")

                    model = str(exif.get("Model", ""))
                    if model:
                        safe_model = "".join(c for c in model if c.isalnum() or c in (" ", "_")).rstrip()
                        if safe_model:
                            dest_folder = os.path.join("Images", safe_model)
                else:
                    dest_folder = "Images"
            except Exception:
                dest_folder = "Images"
        else:
            for cat, exts in self.categories.items():
                if ext in exts:
                    dest_folder = cat
                    break
        return "move", dest_folder, os.path.basename(file_path)

    @Slot(str)
    def organize_single_file(self, path):
        """Organizes a single file (used by watcher). Includes conflict handling."""
        self.file_metadata_cache = {}  # Clear cache for single file
        time.sleep(0.5)
        if not os.path.exists(path) or os.path.basename(path).startswith("."):
            return

        self.log_message.emit(f"Watcher: New file - {os.path.basename(path)}")
        action_type, dest_rel_path, new_name = self.determine_destination_and_action(path)
        base_path = os.path.dirname(path)

        try:
            if action_type == "delete":
                os.remove(path)
                self.log_message.emit(f"  - DELETED '{os.path.basename(path)}' by rule.")
                self.single_file_organized.emit(f"Deleted: {os.path.basename(path)}")
                return  # Stop processing after delete

            # --- Logic for move, copy, rename ---
            log_dest_folder = ""
            final_dest = ""
            dest_path_dir = ""

            if action_type == "move" or action_type == "copy":
                dest_path_dir = os.path.join(base_path, dest_rel_path)
                final_dest = os.path.join(dest_path_dir, new_name)
                log_dest_folder = dest_rel_path
            elif action_type == "rename":
                dest_path_dir = base_path
                final_dest = os.path.join(dest_path_dir, new_name)
                log_dest_folder = "(Renamed)"
            else:
                self.log_message.emit(f"  - SKIPPED: Unknown action type '{action_type}'")
                return

            if not os.path.exists(dest_path_dir):
                os.makedirs(dest_path_dir)

            target_exists = os.path.exists(final_dest)
            is_same_file = os.path.normpath(path) == os.path.normpath(final_dest)

            if target_exists and not is_same_file:
                strategy = self.conflict_strategy
                if strategy == "skip":
                    self.log_message.emit(f"  - SKIPPED (Conflict)")
                    return
                elif strategy == "rename":
                    base, ext = os.path.splitext(new_name)
                    count = 1
                    while os.path.exists(final_dest):
                        final_dest = os.path.join(dest_path_dir, f"{base} ({count}){ext}")
                        count += 1
                    self.log_message.emit(f"  - RENAMED (Conflict) to '{os.path.basename(final_dest)}'")
                # REMOVED: Overwrite logic

            # Execute the action
            if action_type == "move":
                shutil.move(path, final_dest)
                self.log_message.emit(f"  - Moved to '{log_dest_folder}'")
                self.single_file_organized.emit(f"Organized: {os.path.basename(final_dest)}")
            elif action_type == "copy":
                shutil.copy2(path, final_dest)
                self.log_message.emit(f"  - Copied to '{log_dest_folder}'")
                self.single_file_organized.emit(f"Copied: {os.path.basename(final_dest)}")  # copy2 preserves metadata
            elif action_type == "rename":
                if not is_same_file:
                    shutil.move(path, final_dest)
                    self.log_message.emit(f"  - Renamed to '{os.path.basename(final_dest)}'")
                    self.single_file_organized.emit(f"Renamed: {os.path.basename(final_dest)}")
                else:
                    self.log_message.emit(f"  - INFO: Rename rule resulted in same filename, skipped.")

        except Exception as exc:
            self.log_message.emit(f"  - ERROR organizing {os.path.basename(path)}: {exc}")

    @Slot(str, bool, bool)  # Added is_recursive flag
    def run_organization_preview(self, path, backup_first, is_recursive):
        """Generates the preview of organization actions, including copy/delete."""
        self.file_metadata_cache = {}  # Clear cache for new preview
        self.log_message.emit("--- Generating Organization Preview ---")
        if is_recursive:
            self.log_message.emit("Scanning subfolders (recursive)...")
        else:
            self.log_message.emit("Scanning root folder (non-recursive)...")

        items_to_scan = []
        category_folders = [d.lower() for d in self.categories.keys()]

        try:
            for root, dirs, files in os.walk(path, topdown=True):
                # --- Pruning logic ---
                # If we are in the root folder, prune category dirs
                if os.path.normpath(root) == os.path.normpath(path):
                    dirs[:] = [d for d in dirs if d.lower() not in category_folders]
                # If we are *not* in the root folder (i.e., a subfolder)
                else:
                    # Check if this subfolder is a category folder (shouldn't happen if pruning worked, but safe)
                    rel_path = os.path.relpath(root, path)
                    first_level_dir = rel_path.split(os.path.sep)[0]
                    if first_level_dir.lower() in category_folders:
                        dirs[:] = []  # Don't descend further
                        files = []  # Don't process files in this dir

                # --- Add files to list ---
                for f in files:
                    if not f.startswith(".") and f != LOG_FILE_NAME and not f.startswith("~$"):
                        items_to_scan.append(os.path.join(root, f))

                # --- Non-recursive check ---
                if not is_recursive:
                    break  # Stop after processing root folder

        except OSError as exc:
            self.log_message.emit(f"Error scanning folder: {exc}")
            self.finished.emit()
            return

        # --- This point is reached after scanning ---

        total_items = len(items_to_scan)
        if total_items == 0:
            self.log_message.emit("No files to preview.")
            self.finished.emit()
            return

        self.log_message.emit(f"Found {total_items} files. Generating preview...")
        proposed_actions = []
        move_log_temp = {}

        for i, item_path in enumerate(items_to_scan):  # Changed from name to item_path

            action_type, dest_rel_path_or_dir, new_name = self.determine_destination_and_action(item_path)

            final_dest_path = ""
            display_action_detail = ""
            action_status = ""

            if action_type == "delete":
                final_dest_path = item_path  # Use original path for logging deletion
                display_action_detail = "**DELETE**"
                action_status = "(Will be deleted)"

            elif action_type == "move" or action_type == "copy":
                dest_path_dir = os.path.join(path, dest_rel_path_or_dir)
                final_dest_path = os.path.join(dest_path_dir, new_name)
                display_action_detail = os.path.relpath(final_dest_path, start=path)
            elif action_type == "rename":
                # For renames, the destination directory is the file's current directory
                dest_path_dir = os.path.dirname(item_path)
                final_dest_path = os.path.join(dest_path_dir, new_name)
                display_action_detail = os.path.relpath(final_dest_path, start=path)
            else:
                continue  # Skip unknown actions

            if not final_dest_path:
                continue

            is_same_file = os.path.normpath(item_path) == os.path.normpath(final_dest_path)

            if action_type != "delete":  # Conflict handling only for move/copy/rename
                if is_same_file:  # If file is already in the correct place, skip
                    continue

                target_exists = os.path.exists(final_dest_path)
                if target_exists:
                    strategy = self.conflict_strategy
                    if strategy == "skip":
                        action_status = "(Skipped - Exists)"
                        final_dest_path = None
                    elif strategy == "rename":
                        base, ext = os.path.splitext(new_name)
                        count = 1
                        temp_dest = final_dest_path
                        while os.path.exists(temp_dest):
                            temp_dest = os.path.join(dest_path_dir, f"{base} ({count}){ext}")
                            count += 1
                        display_action_detail = os.path.relpath(temp_dest, start=path)
                        action_status = "(Renamed - Conflict)"
                        final_dest_path = temp_dest
                    # REMOVED: Overwrite logic

            if final_dest_path:
                proposed_actions.append((action_type, item_path, f"{display_action_detail} {action_status}"))
                move_log_temp[final_dest_path] = (item_path, action_type)  # Store action type

            # Update progress bar periodically
            if (i + 1) % 100 == 0 or (i + 1) == total_items:
                self.progress_updated.emit(i + 1, total_items)

        self.log_message.emit("--- Preview Generated ---")
        self.organization_preview_ready.emit(proposed_actions, move_log_temp, backup_first)

    @Slot(str, dict, bool)
    def execute_organization_moves(self, path, move_log_data, backup_first):
        """Executes the actual file moves, copies, renames, and deletes after preview approval."""
        self.log_message.emit("--- Starting Organization ---")
        if backup_first:
            # ... (backup logic remains the same) ...
            self.log_message.emit("Creating backup archive...")
            backup_name = f"organizer_backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
            backup_path = os.path.join(os.path.dirname(path), backup_name)
            try:
                with zipfile.ZipFile(backup_path, "w", zipfile.ZIP_DEFLATED) as zipf:
                    for root, dirs, files in os.walk(path):
                        # Exclude log file from backup
                        if LOG_FILE_NAME in files:
                            files.remove(LOG_FILE_NAME)
                        # Exclude empty dirs potentially? (optional)
                        for file in files:
                            file_full_path = os.path.join(root, file)
                            arcname = os.path.relpath(file_full_path, path)
                            zipf.write(file_full_path, arcname)
                self.log_message.emit(f"Backup created: {backup_path}")
            except Exception as exc:
                self.log_message.emit(f"  - ERROR creating backup: {exc}")

        total_items = len(move_log_data)
        items_processed = 0
        actual_move_log = {}

        for final_dest_or_original, (item_path, action_type) in move_log_data.items():
            try:
                if action_type == "delete":
                    os.remove(item_path)
                    actual_move_log[f"deleted_{items_processed}"] = {"original": item_path, "action": "delete"}  # Log deletion
                    self.log_message.emit(f"Deleted '{os.path.relpath(item_path, path)}'")

                elif action_type == "move" or action_type == "rename":
                    dest_path_dir = os.path.dirname(final_dest_or_original)
                    if not os.path.exists(dest_path_dir):
                        os.makedirs(dest_path_dir)
                    shutil.move(item_path, final_dest_or_original)
                    actual_move_log[final_dest_or_original] = {"original": item_path, "action": action_type}  # Log action
                    if action_type == "move":
                        self.log_message.emit(f"Moved '{os.path.relpath(item_path, path)}' -> '{os.path.relpath(final_dest_or_original, path)}'")
                    else:
                        self.log_message.emit(f"Renamed '{os.path.relpath(item_path, path)}' -> '{os.path.relpath(final_dest_or_original, path)}'")

                elif action_type == "copy":
                    dest_path_dir = os.path.dirname(final_dest_or_original)
                    if not os.path.exists(dest_path_dir):
                        os.makedirs(dest_path_dir)
                    shutil.copy2(item_path, final_dest_or_original)  # copy2 preserves metadata
                    actual_move_log[final_dest_or_original] = {"original": item_path, "action": action_type}  # Log action
                    self.log_message.emit(f"Copied '{os.path.relpath(item_path, path)}' -> '{os.path.relpath(final_dest_or_original, path)}'")

            except Exception as exc:
                self.log_message.emit(f"  - ERROR {action_type}ing '{os.path.basename(item_path)}': {exc}")

            items_processed += 1
            self.progress_updated.emit(items_processed, total_items)

        if actual_move_log:
            try:
                # Save log with timestamp (basic history, replace previous)
                # log_filename = f".organizer_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                # For now, keep overwriting the single log file for simple undo
                with open(os.path.join(path, LOG_FILE_NAME), "w") as f:
                    json.dump(actual_move_log, f, indent=4)
            except IOError as exc:
                self.log_message.emit(f"  - ERROR writing log file: {exc}")

        self.log_message.emit("--- Organization Complete! ---")
        self.finished.emit()

    @Slot(str)
    def run_deorganization(self, path):
        self.log_message.emit("\n--- Starting De-organization ---")
        log_path = os.path.join(path, LOG_FILE_NAME)
        if not os.path.exists(log_path):
            self.log_message.emit("ERROR: Log file not found.")
            self.finished.emit()
            return

        try:
            with open(log_path, "r") as f:
                move_log = json.load(f)
            total = len(move_log)
            items_processed = 0

            # --- Iterate in reverse for potentially safer directory handling ---
            organized_paths = list(move_log.keys())

            for key in organized_paths:  # Use key to handle deletion logs
                log_entry = move_log[key]

                # --- Handle potential old log format ---
                if isinstance(log_entry, str):
                    original = log_entry
                    action_type = "move"  # Assume move for old format
                    organized = key  # Key was the destination path
                else:
                    original = log_entry.get("original")
                    action_type = log_entry.get("action", "move")
                    organized = key if action_type != "delete" else None  # Deletion key is arbitrary

                items_processed += 1

                if action_type == "delete":
                    self.log_message.emit(f"  - SKIPPED Revert: Cannot undo deletion of '{os.path.basename(original)}'")
                    continue  # Cannot revert deletion

                if not organized or not original:
                    self.log_message.emit(f"  - WARN: Invalid log entry skipped: {key}")
                    continue

                if os.path.exists(organized):
                    try:
                        original_dir = os.path.dirname(original)
                        if not os.path.exists(original_dir):
                            os.makedirs(original_dir)

                        restored_original_path = original  # Path to restore to

                        if os.path.exists(original):
                            base, ext = os.path.splitext(original)
                            count = 1
                            new_original = os.path.join(original_dir, f"{base} (restored {count}){ext}")
                            while os.path.exists(new_original):
                                count += 1
                                new_original = os.path.join(original_dir, f"{base} (restored {count}){ext}")
                            self.log_message.emit(
                                f"  - WARN: Original exists, renaming restored file to {os.path.basename(new_original)}"
                            )
                            restored_original_path = new_original

                        # Only move back if it was moved or renamed originally
                        if action_type == "move" or action_type == "rename":
                            shutil.move(organized, restored_original_path)
                            self.log_message.emit(f"'{os.path.basename(restored_original_path)}' <- Moved back")
                        elif action_type == "copy":
                            # Optionally delete the copied file? For now, just log.
                            self.log_message.emit(
                                f"  - INFO: Original action was 'copy', leaving '{os.path.basename(organized)}' in place."
                            )

                    except Exception as exc:
                        self.log_message.emit(f"  - ERROR moving back '{os.path.basename(organized)}': {exc}")

                elif action_type != "delete":  # Don't warn if a deleted file isn't found
                    self.log_message.emit(f"  - WARN: Organized file not found, skipping: {organized}")

                self.progress_updated.emit(items_processed, total)

            # --- Cleanup empty directories (same logic as before) ---
            organized_dirs = set()
            for key, log_entry in move_log.items():
                action = log_entry.get("action", "move") if isinstance(log_entry, dict) else "move"
                if action != "delete":
                    organized_dirs.add(os.path.dirname(key))

            for org_dir in organized_dirs:
                try:
                    # Check if exists and is empty AFTER all moves
                    if os.path.exists(org_dir) and not os.listdir(org_dir) and os.path.normpath(org_dir) != os.path.normpath(path):
                        os.rmdir(org_dir)
                        self.log_message.emit(f"Removed empty directory: {os.path.basename(org_dir)}")
                except OSError as exc:
                    self.log_message.emit(f"  - WARN: Could not remove dir {os.path.basename(org_dir)}: {exc}")

            os.remove(log_path)
            self.log_message.emit("Removed log file.")
            self.log_message.emit("--- De-organization Complete! ---")

        except json.JSONDecodeError:
            self.log_message.emit(f"ERROR: Could not read log file '{LOG_FILE_NAME}'.")
        except Exception as exc:
            self.log_message.emit(f"An unexpected error occurred: {exc}")
        finally:
            self.finished.emit()

    @Slot(str)
    def run_duplicate_scan(self, path):
        # ... (duplicate scan logic remains the same) ...
        self.log_message.emit("--- Starting Duplicate File Scan ---")
        hashes = {}
        file_list = []
        for root, _, files in os.walk(path):
            for file in files:
                if not file.startswith(".") and file != LOG_FILE_NAME:
                    file_list.append(os.path.join(root, file))
        total = len(file_list)
        if total == 0:
            self.log_message.emit("No files found to scan.")
            self.finished.emit()
            return
        processed_files = 0
        for filepath in file_list:
            try:
                hasher = hashlib.md5()
                with open(filepath, "rb") as f:
                    while chunk := f.read(8192 * 16):
                        hasher.update(chunk)
                file_hash = hasher.hexdigest()
                if file_hash in hashes:
                    hashes[file_hash].append(filepath)
                else:
                    hashes[file_hash] = [filepath]
            except (IOError, OSError) as exc:
                self.log_message.emit(f"Could not read {os.path.basename(filepath)}: {exc}")
            processed_files += 1
            self.progress_updated.emit(processed_files, total)
        duplicates = {hash_val: files for hash_val, files in hashes.items() if len(files) > 1}
        self.log_message.emit(f"--- Duplicate Scan Complete: Found {len(duplicates)} sets of duplicates. ---")
        self.duplicate_scan_finished.emit(duplicates)

    @Slot(str)
    def scan_folder_stats(self, path):
        # ... (folder stats logic remains the same) ...
        """Scans folder for size and file count."""
        count = 0
        size_bytes = 0
        type_counts = {}
        try:
            for root, _, files in os.walk(path):
                for file in files:
                    if not file.startswith(".") and file != LOG_FILE_NAME:
                        count += 1
                        filepath = os.path.join(root, file)
                        try:
                            size_bytes += os.path.getsize(filepath)
                        except OSError:
                            pass
                        ext = os.path.splitext(file)[1].lower()
                        if ext:
                            type_counts[ext] = type_counts.get(ext, 0) + 1

            # Calculate size string here to avoid int overflow in signal
            size_mb = size_bytes / (1024 * 1024)
            size_gb = size_bytes / (1024 * 1024 * 1024)
            size_str = f"{size_gb:.2f} GB" if size_gb >= 1 else f"{size_mb:.2f} MB"
            self.folder_stats_ready.emit(count, size_str, type_counts)  # Pass string

        except Exception as exc:
            self.log_message.emit(f"Error scanning folder stats: {exc}")
        finally:
            self.finished.emit()  # Moved finished emit here

    @Slot(str)
    def run_empty_folder_cleanup(self, path):
        # ... (cleanup logic remains the same) ...
        """Finds and deletes empty subfolders."""
        self.log_message.emit("--- Scanning for empty folders... ---")
        folders_removed = 0
        try:
            # Walk bottom-up to remove child folders first
            for root, dirs, files in os.walk(path, topdown=False):
                # Don't delete the root folder itself
                if os.path.normpath(root) == os.path.normpath(path):
                    continue

                if not dirs and not files:
                    try:
                        os.rmdir(root)
                        self.log_message.emit(f"Removed empty folder: {os.path.relpath(root, path)}")
                        folders_removed += 1
                    except OSError as exc:
                        self.log_message.emit(f"  - WARN: Could not remove {os.path.relpath(root, path)}: {exc}")

            if folders_removed == 0:
                self.log_message.emit("No empty folders found.")
            else:
                self.log_message.emit(f"--- Cleanup complete: {folders_removed} empty folders removed. ---")
        except Exception as exc:
            self.log_message.emit(f"  - ERROR during cleanup: {exc}")
        finally:
            self.finished.emit()


if WATCHDOG_AVAILABLE:

    class FileSystemWatcher(QObject):
        file_created = Signal(str)

        class EventHandler(FileSystemEventHandler):
            def __init__(self, signal_emitter):
                super().__init__()
                self.signal_emitter = signal_emitter

            def on_created(self, event):
                time.sleep(1)  # Ensure file is fully written
                if not event.is_directory:
                    try:
                        self.signal_emitter.emit(event.src_path)
                    except Exception:
                        # avoid crashing the watchdog thread on unexpected errors
                        pass

        def __init__(self, path):
            super().__init__()
            self.path = path
            self.observer = Observer()
            event_handler = self.EventHandler(self.file_created)
            self.observer.schedule(event_handler, self.path, recursive=False)

        def start(self):
            self.observer.start()

        def stop(self):
            self.observer.stop()
            self.observer.join()


else:

    class FileSystemWatcher(QObject):
        file_created = Signal(str)

        def __init__(self, path):
            super().__init__()
            print("Watchdog library not installed.")

        def start(self):
            pass

        def stop(self):
            pass


class FileOrganizerWindow(QMainWindow):
    # Updated signal to include recursive flag
    start_organization_preview_signal = Signal(str, bool, bool)
    execute_organization_signal = Signal(str, dict, bool)
    start_deorganization_signal = Signal(str)
    start_duplicate_scan_signal = Signal(str)
    organize_one_file_signal = Signal(str)
    scan_folder_stats_signal = Signal(str)
    start_cleanup_signal = Signal(str)  # Signal for empty folder cleanup

    def __init__(self):
        super().__init__()
        self.setObjectName("MainWindow")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground, True)
        self.setGeometry(100, 100, 900, 600)
        self.setAcceptDrops(True)
        self.watcher = None
        self.watcher_thread = None
        self.current_theme = "dark"
        self.current_preview_data = {}
        self.central_widget = QWidget(self)
        self.central_widget.setObjectName("CentralWidget")
        self.setCentralWidget(self.central_widget)
        main_layout = QVBoxLayout(self.central_widget)
        main_layout.setContentsMargins(1, 1, 1, 1)
        main_layout.setSpacing(0)
        content_card = QWidget()
        content_card.setObjectName("ContentCard")
        main_layout.addWidget(content_card)
        card_layout = QVBoxLayout(content_card)
        card_layout.setContentsMargins(0, 0, 0, 0)
        card_layout.setSpacing(0)
        self.create_title_bar()
        card_layout.addWidget(self.title_bar)
        self.create_toolbar()
        card_layout.addWidget(self.toolbar)
        self.create_content_area()
        card_layout.addWidget(self.content_area, 1)
        self.setup_worker_thread()
        self.setup_tray_icon()
        self.load_app_settings()

    def create_toolbar(self):
        self.toolbar = QWidget()
        self.toolbar.setObjectName("ToolBar")

        # Use three rows for the toolbar
        toolbar_layout_main = QVBoxLayout(self.toolbar)
        toolbar_layout_main.setContentsMargins(10, 5, 10, 5)
        toolbar_layout_main.setSpacing(8)  # Spacing between rows

        # --- First Row: Path and Browse ---
        toolbar_layout_row1 = QHBoxLayout()
        toolbar_layout_row1.setSpacing(8)
        toolbar_layout_row1.addWidget(QLabel("Folder:"))
        self.path_input = QLineEdit()
        self.path_input.setReadOnly(True)
        self.path_input.setPlaceholderText("Select or drop folder...")
        toolbar_layout_row1.addWidget(self.path_input, 1)  # Give path input stretch priority
        self.browse_btn = QPushButton(qta.icon("fa5s.folder-open", color="#cbd5e1"), "")
        self.browse_btn.setToolTip("Browse for folder")
        self.browse_btn.clicked.connect(self.select_folder)
        toolbar_layout_row1.addWidget(self.browse_btn)

        toolbar_layout_main.addLayout(toolbar_layout_row1)

        # --- Second Row: Main Actions ---
        toolbar_layout_row2 = QHBoxLayout()
        toolbar_layout_row2.setSpacing(8)
        self.organize_btn = QPushButton(qta.icon("fa5s.magic", color="#111111"), " Organize")
        self.organize_btn.setObjectName("OrganizeBtn")
        self.organize_btn.setToolTip("Organize selected folder (shows preview)")
        self.organize_btn.clicked.connect(self.start_organization_preview)
        toolbar_layout_row2.addWidget(self.organize_btn)
        self.deorganize_btn = QPushButton(qta.icon("fa5s.undo", color="#e0e0e0"), " De-organize")
        self.deorganize_btn.setToolTip("Undo last organization")
        self.deorganize_btn.clicked.connect(self.start_deorganization)
        toolbar_layout_row2.addWidget(self.deorganize_btn)
        self.duplicates_btn = QPushButton(qta.icon("fa5s.copy", color="#e0e0e0"), " Duplicates")
        self.duplicates_btn.setToolTip("Find duplicate files")
        self.duplicates_btn.clicked.connect(self.start_duplicate_scan)
        toolbar_layout_row2.addWidget(self.duplicates_btn)
        self.cleanup_btn = QPushButton(qta.icon("fa5s.broom", color="#e0e0e0"), " Clean Empty")
        self.cleanup_btn.setToolTip("Delete empty subfolders")
        self.cleanup_btn.clicked.connect(self.start_empty_folder_cleanup)
        toolbar_layout_row2.addWidget(self.cleanup_btn)
        toolbar_layout_row2.addStretch()  # Add stretch at the end of actions

        toolbar_layout_main.addLayout(toolbar_layout_row2)

        # --- Third Row: Settings and Toggles ---
        toolbar_layout_row3 = QHBoxLayout()
        toolbar_layout_row3.setSpacing(8)

        # --- Profile/Category Buttons ---
        self.load_profile_btn = QPushButton(qta.icon("fa5s.folder-open", color="#e0e0e0"), " Load Profile")
        self.load_profile_btn.setToolTip("Load rules and settings from a profile")
        self.load_profile_btn.clicked.connect(self.load_profile)
        toolbar_layout_row3.addWidget(self.load_profile_btn)
        self.save_profile_btn = QPushButton(qta.icon("fa5s.save", color="#e0e0e0"), " Save Profile As...")
        self.save_profile_btn.setToolTip("Save current rules and settings to a profile")
        self.save_profile_btn.clicked.connect(self.save_profile_as)
        toolbar_layout_row3.addWidget(self.save_profile_btn)
        self.categories_btn = QPushButton(qta.icon("fa5s.stream", color="#e0e0e0"), " Categories")
        self.categories_btn.setToolTip("Edit file categories")
        self.categories_btn.clicked.connect(self.open_category_editor)
        toolbar_layout_row3.addWidget(self.categories_btn)

        toolbar_layout_row3.addStretch()  # Stretch to separate settings buttons from toggles

        # --- Toggles and Options ---
        toolbar_layout_row3.addWidget(QLabel("On Conflict:"))
        self.conflict_combo = QComboBox()
        self.conflict_combo.addItems(["Rename", "Skip"])  # Removed "Overwrite"
        self.conflict_combo.setToolTip("Action to take if file already exists")
        self.conflict_combo.currentTextChanged.connect(self.conflict_strategy_changed)
        toolbar_layout_row3.addWidget(self.conflict_combo)

        sep2 = QFrame()
        sep2.setFrameShape(QFrame.Shape.VLine)
        sep2.setFrameShadow(QFrame.Shadow.Sunken)
        toolbar_layout_row3.addWidget(sep2)

        # FIX: Connect toggled signals correctly to save settings without passing bool arg
        self.recursive_check = QCheckBox("Scan Subfolders")
        self.recursive_check.setToolTip("Scan all subfolders recursively")
        self.recursive_check.toggled.connect(lambda: self.save_app_settings())
        toolbar_layout_row3.addWidget(self.recursive_check)
        self.backup_check = QCheckBox("Backup First")
        self.backup_check.setToolTip("Create a ZIP backup before organizing")
        self.backup_check.toggled.connect(lambda: self.save_app_settings())
        toolbar_layout_row3.addWidget(self.backup_check)

        sep3 = QFrame()
        sep3.setFrameShape(QFrame.Shape.VLine)
        sep3.setFrameShadow(QFrame.Shadow.Sunken)
        toolbar_layout_row3.addWidget(sep3)

        self.watcher_check = QCheckBox()
        self.watcher_check.setObjectName("WatcherToggle")
        self.watcher_check.setToolTip("Automatically organize new files")
        if not WATCHDOG_AVAILABLE:
            self.watcher_check.setEnabled(False)
            self.watcher_check.setToolTip("Install 'watchdog' library to enable watcher")
        self.watcher_check.toggled.connect(self.toggle_watcher)
        toolbar_layout_row3.addWidget(self.watcher_check)
        toolbar_layout_row3.addWidget(QLabel("Watch"))

        toolbar_layout_main.addLayout(toolbar_layout_row3)

    def create_content_area(self):
        self.content_area = QWidget()
        self.content_area.setObjectName("ContentArea")
        content_layout = QVBoxLayout(self.content_area)
        content_layout.setContentsMargins(20, 10, 20, 15)
        content_layout.setSpacing(10)
        self.stats_label = QLabel("Select a folder to see stats.")
        self.stats_label.setObjectName("StatsLabel")
        content_layout.addWidget(self.stats_label)

        # --- Log Tools (Export Button) ---
        log_tools_widget = QWidget()
        log_tools_widget.setObjectName("LogToolsWidget")
        log_tools_layout = QHBoxLayout(log_tools_widget)
        log_tools_layout.setContentsMargins(0, 0, 0, 5)
        log_tools_layout.setSpacing(10)
        log_tools_layout.addStretch()
        self.export_log_btn = QPushButton(qta.icon("fa5s.file-export"), " Export Log")
        self.export_log_btn.clicked.connect(self.export_log)
        log_tools_layout.addWidget(self.export_log_btn)
        content_layout.addWidget(log_tools_widget)

        self.log_area = QTextEdit()
        self.log_area.setReadOnly(True)
        content_layout.addWidget(self.log_area, 1)
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        content_layout.addWidget(self.progress_bar)

    def create_title_bar(self):
        self.title_bar = QWidget(self)
        self.title_bar.setObjectName("TitleBar")
        self.title_bar.setFixedHeight(40)
        title_layout = QHBoxLayout(self.title_bar)
        title_layout.setContentsMargins(15, 0, 5, 0)
        title_icon = QLabel()
        title_icon.setPixmap(qta.icon("fa5s.rocket", color="#007AFF", scale_factor=1.0).pixmap(20, 20))
        title_layout.addWidget(title_icon)
        title_label = QLabel("File Organizer")
        title_label.setObjectName("TitleLabel")
        title_layout.addWidget(title_label)
        title_layout.addStretch()
        self.help_btn = QPushButton("")
        self.help_btn.setObjectName("HelpBtn")
        self.help_btn.setIcon(qta.icon("fa5s.question-circle", color=self._get_icon_color()))
        self.help_btn.setFixedSize(28, 28)
        self.help_btn.setToolTip("How to use")
        self.help_btn.clicked.connect(self.show_manual)
        title_layout.addWidget(self.help_btn)
        self.theme_toggle_btn = QPushButton("")
        self.theme_toggle_btn.setObjectName("ThemeToggleBtn")
        self.theme_toggle_btn.setFixedSize(28, 28)
        self.theme_toggle_btn.clicked.connect(self.toggle_theme)
        title_layout.addWidget(self.theme_toggle_btn)
        btn_size = 26
        minimize_btn = QPushButton(qta.icon("fa5.window-minimize", color="#e0e0e0"), "")
        minimize_btn.setObjectName("WindowButton")
        minimize_btn.setFixedSize(btn_size, btn_size)
        minimize_btn.clicked.connect(self.showMinimized)
        self.maximize_btn = QPushButton(qta.icon("fa5.window-maximize", color="#e0e0e0"), "")
        self.maximize_btn.setObjectName("WindowButton")
        self.maximize_btn.setFixedSize(btn_size, btn_size)
        self.maximize_btn.clicked.connect(self.toggle_maximize)
        close_btn = QPushButton(qta.icon("fa5s.times", color="#e0e0e0"), "")
        close_btn.setObjectName("WindowButton")
        close_btn.setProperty("id", "CloseButton")
        close_btn.setFixedSize(btn_size, btn_size)
        close_btn.clicked.connect(self.close)
        title_layout.addWidget(minimize_btn)
        title_layout.addWidget(self.maximize_btn)
        title_layout.addWidget(close_btn)
        self.oldPos = self.pos()

    def dragEnterEvent(self, event: QDragEnterEvent):
        mime_data = event.mimeData()
        if mime_data.hasUrls() and len(mime_data.urls()) == 1:
            url = mime_data.urls()[0]
            if url.isLocalFile() and os.path.isdir(url.toLocalFile()):
                event.acceptProposedAction()

    def dropEvent(self, event: QDropEvent):
        url = event.mimeData().urls()[0]
        folder_path = url.toLocalFile()
        self.path_input.setText(folder_path)
        self.log_area.clear()
        self.log_area.append(f"Selected folder: {folder_path}")
        if self.watcher and self.watcher.path != folder_path:
            self.watcher_check.setChecked(False)
        self.scan_folder_stats_signal.emit(folder_path)

    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton and event.pos().y() < self.title_bar.height():
            self.oldPos = event.globalPosition().toPoint()

    def mouseMoveEvent(self, event):
        if event.buttons() == Qt.MouseButton.LeftButton and event.pos().y() < self.title_bar.height():
            delta = QPoint(event.globalPosition().toPoint() - self.oldPos)
            self.move(self.x() + delta.x(), self.y() + delta.y())
            self.oldPos = event.globalPosition().toPoint()

    def toggle_maximize(self):
        if self.isMaximized():
            self.showNormal()
            self.maximize_btn.setIcon(qta.icon("fa5.window-maximize", color=self._get_icon_color()))
        else:
            self.showMaximized()
            self.maximize_btn.setIcon(qta.icon("fa5.window-restore", color=self._get_icon_color()))

    def setup_worker_thread(self):
        self.thread = QThread()
        self.worker = Worker()
        self.worker.moveToThread(self.thread)
        self.worker.log_message.connect(self.update_log)
        self.worker.progress_updated.connect(self.update_progress)
        self.worker.finished.connect(self.on_task_finished)
        self.worker.single_file_organized.connect(self.show_tray_notification)
        self.worker.duplicate_scan_finished.connect(self.on_duplicate_scan_finished)
        self.worker.organization_preview_ready.connect(self.on_organization_preview_ready)
        self.worker.folder_stats_ready.connect(self.on_folder_stats_ready)

        # Connect updated/new signals
        self.start_organization_preview_signal.connect(self.worker.run_organization_preview)
        self.execute_organization_signal.connect(self.worker.execute_organization_moves)
        self.start_deorganization_signal.connect(self.worker.run_deorganization)
        self.organize_one_file_signal.connect(self.worker.organize_single_file)
        self.start_duplicate_scan_signal.connect(self.worker.run_duplicate_scan)
        self.scan_folder_stats_signal.connect(self.worker.scan_folder_stats)
        self.start_cleanup_signal.connect(self.worker.run_empty_folder_cleanup)  # Connect cleanup signal

        self.thread.start()

    def load_app_settings(self, settings_path=None):
        """Loads settings from a given path or the default."""
        if settings_path is None:
            settings_path = SETTINGS_FILE_NAME

        try:
            with open(settings_path, "r") as f:
                settings = json.load(f)

            # Load theme only from default settings
            if settings_path == SETTINGS_FILE_NAME:
                self.current_theme = settings.get("theme", "dark")

            # Load other settings from the profile
            saved_strategy = settings.get("conflict_strategy", "rename")  # default to rename
            if saved_strategy == "overwrite":
                saved_strategy = "rename"  # force rename
            self.conflict_combo.setCurrentText(saved_strategy.capitalize())

            self.backup_check.setChecked(settings.get("backup_before_organize", False))
            self.recursive_check.setChecked(settings.get("recursive_scan", False))

            # Pass rules/categories directly to worker
            # Use worker's _get_defaults method if needed
            self.worker.categories = settings.get("categories", self.worker._get_defaults("categories"))
            self.worker.rules = settings.get("rules", self.worker._get_defaults("rules"))
            # ADDED: Update worker's has_content_rules flag
            self.worker.has_content_rules = any(
                cond.get("type") == "content" for rule in self.worker.rules for cond in rule.get("conditions", [])
            )
            self.worker.conflict_strategy = self.conflict_combo.currentText().lower()

            self.log_area.append(f"--- Settings loaded from {os.path.basename(settings_path)} ---")

        except (FileNotFoundError, json.JSONDecodeError) as exc:
            self.log_area.append(f"--- Could not load settings from {os.path.basename(settings_path)}: {exc} ---")
            if settings_path == SETTINGS_FILE_NAME:  # Only apply defaults if default settings fail
                self.log_area.append("--- Loading default settings ---")
                self.current_theme = "dark"
                self.conflict_combo.setCurrentText("Rename")
                self.backup_check.setChecked(False)
                self.recursive_check.setChecked(False)
                # FIX: Call worker's load_settings to properly load defaults
                self.worker.load_settings()

        self.apply_theme()  # Apply theme after loading

    def save_app_settings(self, settings_path=None):
        """Saves settings to a given path or the default."""
        is_default_save = settings_path is None
        if is_default_save:
            settings_path = SETTINGS_FILE_NAME

        settings_to_save = {}

        # Only save theme to default file
        if is_default_save:
            settings_to_save["theme"] = self.current_theme

        # Save profile-specific settings
        settings_to_save["conflict_strategy"] = self.conflict_combo.currentText().lower()
        settings_to_save["backup_before_organize"] = self.backup_check.isChecked()
        settings_to_save["recursive_scan"] = self.recursive_check.isChecked()
        settings_to_save["rules"] = self.worker.rules
        settings_to_save["categories"] = self.worker.categories

        try:
            # FIX: Ensure settings_path is a string before using os.path.basename
            log_filename = os.path.basename(str(settings_path)) if settings_path else os.path.basename(SETTINGS_FILE_NAME)
            with open(settings_path, "w") as f:
                json.dump(settings_to_save, f, indent=4)
            self.log_area.append(f"--- Settings saved to {log_filename} ---")
        except IOError as exc:
            self.log_area.append(f"--- ERROR saving settings: {exc} ---")
            self.show_error_message(f"Error saving settings: {exc}")

        # Ensure worker is in sync, especially if saving the default file
        if is_default_save:
            self.worker.load_settings()

    # --- Profile Functions ---
    def load_profile(self):
        """Loads rules, categories, and settings from a .json file."""
        file_path, _ = QFileDialog.getOpenFileName(self, "Load Profile", "", "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.load_app_settings(file_path)

    def save_profile_as(self):
        """Saves current rules, categories, and settings to a new .json file."""
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Profile As", "organizer_profile.json", "JSON Files (*.json);;All Files (*)")
        if file_path:
            self.save_app_settings(file_path)

    def _get_icon_color(self):
        return "#374151" if self.current_theme == "light" else "#e0e0e0"

    def toggle_theme(self):
        self.current_theme = "light" if self.current_theme == "dark" else "dark"
        self.apply_theme()
        self.save_app_settings()  # Save theme to default file

    def apply_theme(self):
        icon_color = self._get_icon_color()
        if self.current_theme == "light":
            self.setStyleSheet(LIGHT_BLUE_THEME)
            self.theme_toggle_btn.setIcon(qta.icon("fa5s.moon", color=icon_color))
        else:
            self.setStyleSheet(DARK_BLUE_THEME)
            self.theme_toggle_btn.setIcon(qta.icon("fa5s.sun", color=icon_color))

        # Update title bar icons
        for btn in self.title_bar.findChildren(QPushButton):
            object_name = btn.objectName()
            if "minimize" in object_name:
                btn.setIcon(qta.icon("fa5.window-minimize", color=icon_color))
            elif "maximize" in object_name:
                btn.setIcon(qta.icon("fa5.window-maximize" if not self.isMaximized() else "fa5.window-restore", color=icon_color))
            elif "CloseButton" in object_name:
                btn.setIcon(qta.icon("fa5s.times", color=icon_color))
            elif "HelpBtn" in object_name:
                btn.setIcon(qta.icon("fa5s.question-circle", color=icon_color))
            elif "ThemeToggleBtn" in object_name:
                btn.setIcon(qta.icon("fa5s.moon" if self.current_theme == "light" else "fa5s.sun", color=icon_color))

        # Update other icons
        self.browse_btn.setIcon(qta.icon("fa5s.folder-open", color=icon_color))
        self.deorganize_btn.setIcon(qta.icon("fa5s.undo", color=icon_color))
        self.duplicates_btn.setIcon(qta.icon("fa5s.copy", color=icon_color))
        self.categories_btn.setIcon(qta.icon("fa5s.stream", color=icon_color))
        self.cleanup_btn.setIcon(qta.icon("fa5s.broom", color=icon_color))
        self.export_log_btn.setIcon(qta.icon("fa5s.file-export", color=icon_color))
        self.load_profile_btn.setIcon(qta.icon("fa5s.folder-open", color=icon_color))
        self.save_profile_btn.setIcon(qta.icon("fa5s.save", color=icon_color))

    def show_error_message(self, text):
        QApplication.beep()
        QMessageBox.warning(self, "Warning", text)

    def select_folder(self):
        if self.watcher:
            self.watcher_check.setChecked(False)  # Stop watcher
        folder = QFileDialog.getExistingDirectory(self, "Select Folder")
        if folder:
            self.path_input.setText(folder)
            self.log_area.clear()
            self.log_area.append(f"Selected folder: {folder}")
            self.stats_label.setText("Scanning folder...")  # Indicate scanning
            self.scan_folder_stats_signal.emit(folder)  # Trigger scan

    def start_organization_preview(self):  # Renamed function
        """Starts the preview generation."""
        path = self.path_input.text()
        if not path:
            self.show_error_message("Please select a folder first.")
            return
        self.set_buttons_enabled(False)
        self.log_area.clear()
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        backup = self.backup_check.isChecked()
        recursive = self.recursive_check.isChecked()  # Get recursive flag
        self.start_organization_preview_signal.emit(path, backup, recursive)  # Emit new signal

    def start_deorganization(self):
        path = self.path_input.text()
        if not path:
            self.show_error_message("Please select a folder first.")
            return
        reply = QMessageBox.question(self, "Confirm", "Are you sure you want to revert?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.set_buttons_enabled(False)
            self.log_area.clear()
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            self.start_deorganization_signal.emit(path)

    def start_duplicate_scan(self):
        path = self.path_input.text()
        if not path:
            self.show_error_message("Please select a folder to scan.")
            return
        if self.watcher_check.isChecked():
            self.show_error_message("Stop watching folder before scanning duplicates.")
            return
        self.set_buttons_enabled(False)
        self.log_area.clear()
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        self.start_duplicate_scan_signal.emit(path)

    def start_empty_folder_cleanup(self):
        """Triggers the empty folder cleanup worker."""
        path = self.path_input.text()
        if not path:
            self.show_error_message("Please select a folder to scan.")
            return
        if self.watcher_check.isChecked():
            self.show_error_message("Stop watching folder before cleaning.")
            return
        reply = QMessageBox.question(self, "Confirm", "Are you sure you want to delete all empty subfolders?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.set_buttons_enabled(False)
            self.log_area.clear()
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            self.start_cleanup_signal.emit(path)

    def export_log(self):
        """Saves the current log content to a text file."""
        log_content = self.log_area.toPlainText()
        if not log_content:
            self.show_error_message("Log is empty.")
            return

        default_name = f"organizer_log_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        file_path, _ = QFileDialog.getSaveFileName(self, "Save Log", default_name, "Text Files (*.txt);;All Files (*)")

        if file_path:
            try:
                with open(file_path, "w", encoding="utf-8") as f:
                    f.write(log_content)
                self.log_area.append(f"--- Log exported to {file_path} ---")
            except IOError as exc:
                self.show_error_message(f"Error exporting log: {exc}")

    @Slot(str)
    def update_log(self, message):
        self.log_area.append(message)

    @Slot(int, int)
    def update_progress(self, current, total):
        if self.progress_bar.maximum() == 0:
            self.progress_bar.setRange(0, total)
        self.progress_bar.setValue(current)

    @Slot()
    def on_task_finished(self):
        if not (self.watcher_check.isEnabled() and self.watcher_check.isChecked()):
            self.set_buttons_enabled(True)
        self.progress_bar.setVisible(False)
        self.progress_bar.setValue(0)

    @Slot(list, dict, bool)
    def on_organization_preview_ready(self, proposed_moves, move_log_data, backup_flag):
        """Shows the preview dialog."""
        self.set_buttons_enabled(True)
        self.progress_bar.setVisible(False)
        self.progress_bar.setValue(0)
        if not proposed_moves:
            QMessageBox.information(self, "Preview", "No files need organizing.")
            return
        dialog = PreviewDialog(proposed_moves, self)  # Pass self (main window) as parent
        if dialog.exec():
            self.log_area.append("--- User approved. Starting organization... ---")
            self.set_buttons_enabled(False)
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            path = self.path_input.text()
            self.execute_organization_signal.emit(path, move_log_data, backup_flag)
            QApplication.beep()
            self.show_tray_notification("Organization started...")
        else:
            self.log_area.append("--- Organization cancelled. ---")
            self.on_task_finished()

    @Slot(int, str, dict)  # MODIFIED: int, str, dict
    def on_folder_stats_ready(self, count, size_str, type_counts):  # MODIFIED: size_str
        """Updates the stats label."""
        self.stats_label.setText(f"Files: {count} | Total Size: {size_str}")
        self.on_task_finished()

    @Slot(dict)
    def on_duplicate_scan_finished(self, duplicates):
        deleted_count = 0
        if not duplicates:
            QMessageBox.information(self, "No Duplicates", "No duplicate files found.")
        else:
            dialog = DuplicateFilesDialog(duplicates, self)
            if dialog.exec():
                deleted_count = dialog.deleted_count
                self.log_area.append(f"--- Duplicate deletion complete ({deleted_count} files removed). ---")
            else:
                self.log_area.append("--- Duplicate deletion cancelled. ---")
        QApplication.beep()
        self.show_tray_notification(f"Duplicate scan finished. {deleted_count} files deleted.")
        self.on_task_finished()

    def set_buttons_enabled(self, enabled):
        buttons = [
            self.browse_btn,
            self.organize_btn,
            self.deorganize_btn,
            self.categories_btn,
            self.theme_toggle_btn,
            self.duplicates_btn,
            self.cleanup_btn,
            self.conflict_combo,
            self.watcher_check,
            self.help_btn,
            self.backup_check,
            self.recursive_check,
            self.export_log_btn,
            self.load_profile_btn,
            self.save_profile_btn,  # Add profile buttons
        ]
        for w in buttons:
            w.setEnabled(enabled)

        watcher_enabled = enabled and WATCHDOG_AVAILABLE and bool(self.path_input.text())
        self.watcher_check.setEnabled(watcher_enabled)

    def open_category_editor(self):
        """Opens the new Category Editor dialog."""
        # Ensure worker has current categories before opening editor
        if not hasattr(self.worker, "categories"):
            self.worker.load_settings()
        dialog = CategoryEditor(self.worker.categories, self)
        if dialog.exec():
            # Save categories directly to worker and then to settings
            self.worker.categories = dialog.get_categories()
            self.save_app_settings()
            self.log_area.append("Categories updated.")

    def conflict_strategy_changed(self, text):
        self.worker.conflict_strategy = text.lower()
        self.save_app_settings()

    def toggle_watcher(self, checked):
        path = self.path_input.text()
        if checked and not path:
            self.show_error_message("Select a folder to watch first.")
            self.watcher_check.setChecked(False)
            return

        # Disable almost all controls when watching
        controls_to_disable = [
            self.conflict_combo,
            self.organize_btn,
            self.deorganize_btn,
            self.duplicates_btn,
            self.browse_btn,
            self.categories_btn,
            self.backup_check,
            self.recursive_check,
            self.cleanup_btn,
            self.load_profile_btn,
            self.save_profile_btn,
        ]
        for control in controls_to_disable:
            control.setEnabled(not checked)

        if checked:
            self.watcher = FileSystemWatcher(path)
            self.watcher_thread = QThread()
            self.watcher.moveToThread(self.watcher_thread)
            # connect the watcher's signal to the worker signal emitter
            self.watcher.file_created.connect(self.organize_one_file_signal.emit)
            self.watcher_thread.started.connect(self.watcher.start)
            self.watcher_thread.start()
            self.log_area.append(f"--- Started watching: {path} ---")
            self.watcher_check.setEnabled(True)
            self.progress_bar.setVisible(True)
            self.progress_bar.setRange(0, 0)
            self.show_tray_notification(f"Started watching: {os.path.basename(path)}")
        elif self.watcher:
            self.watcher.stop()
            self.watcher_thread.quit()
            self.watcher_thread.wait()
            self.watcher = None
            self.watcher_thread = None
            self.log_area.append(f"--- Stopped watching folder ---")
            self.set_buttons_enabled(True)
            self.progress_bar.setVisible(False)
            self.show_tray_notification("Stopped watching folder.")

    def setup_tray_icon(self):
        self.tray_icon = QSystemTrayIcon(qta.icon("fa5s.folder", color="white"), self)
        show_action = QAction("Show", self)
        quit_action = QAction("Quit", self)
        show_action.triggered.connect(self.showNormal)
        quit_action.triggered.connect(self.quit_app)
        tray_menu = QMenu()
        tray_menu.addAction(show_action)
        tray_menu.addAction(quit_action)
        self.tray_icon.setContextMenu(tray_menu)
        self.tray_icon.show()

    def quit_app(self):
        if self.watcher:
            self.watcher.stop()
            self.watcher_thread.quit()
            self.watcher_thread.wait()
        self.save_app_settings()  # Save settings on quit
        QApplication.instance().quit()

    def closeEvent(self, event):
        event.ignore()
        self.hide()
        self.show_tray_notification("Minimized to tray.")

    @Slot(str)
    def show_tray_notification(self, message):
        self.tray_icon.showMessage("File Organizer", message, QIcon(), 3000)

    def show_manual(self):
        """Displays a detailed help message explaining all features."""
        manual_text = """
        <h2 style='margin-bottom: 5px;'>File Organizer Help</h2>

        <b style='font-size: 11pt;'>Core Workflow:</b><br>
        1. Select the folder you want to organize.<br>
        2. Configure categories and options (like backup or recursive scan).<br>
        3. Click 'Organize', review the preview, and approve.<br>
        4. Use other tools like 'Duplicates' or 'Clean Empty' as needed.<br><br>

        <b style='font-size: 11pt;'>Toolbar (Top Row):</b><br>
        <ul>
            <li><b>Folder Path:</b> Displays the currently selected folder. Not directly editable.</li>
            <li><b>Browse Button (Folder Icon):</b> Opens a dialog to select the target folder.</li>
        </ul>
        
        <b style='font-size: 11pt;'>Toolbar (Middle Row):</b><br>
        <ul>
            <li><b>Organize Button (Magic Wand):</b> Starts the organization process based on current rules and settings.
                <ul><li>This will first generate a <b>Preview</b> showing all proposed actions (move, copy, rename, delete).</li>
                    <li>You <u>must</u> approve the preview for any changes to be made.</li></ul>
            </li>
            <li><b>De-organize Button (Undo Icon):</b> Reverts the <u>last</u> 'Organize' operation performed on the currently selected folder by reading the `.organizer_log.json` file.
                <ul><li>It moves files back to their original locations recorded in the log.</li>
                    <li><b>Important:</b> It cannot undo 'delete' or 'copy' actions and might rename files if the original location already has a file with the same name.</li></ul>
            </li>
            <li><b>Duplicates Button (Copy Icon):</b> Scans the selected folder (and subfolders) for files with identical content (using MD5 hash).
                <ul><li>Opens a dialog listing duplicate sets.</li>
                    <li>Allows you to select and permanently delete redundant copies (keeps one original automatically).</li></ul>
            </li>
            <li><b>Clean Empty Button (Broom Icon):</b> Scans the selected folder (and subfolders) and permanently deletes any directories that contain no files or other directories.</li>
        </ul>

        <b style='font-size: 11pt;'>Toolbar (Bottom Row):</b><br>
        <ul>
            <li><b>Load Profile (Folder Icon):</b> Loads Rules, Categories, and organization settings (Conflict, Backup, Recursive) from a saved `.json` profile file.</li>
            <li><b>Save Profile As... (Save Icon):</b> Saves the current Rules, Categories, and settings to a `.json` profile file for later use.</li>
            <li><b>Categories Button (List Icon):</b> Opens the <b>Category Editor</b>.
                <ul><li>Manage the file extensions associated with default sorting folders (Images, Documents, etc.).</li>
                    <li>Add new categories or remove custom ones. The 'Other' category is default and cannot be removed.</li></ul>
            </li>
            <li><b>On Conflict Dropdown:</b> Determines what happens if a file action results in a name collision at the destination:
                <ul><li><b>Rename:</b> Appends a number (e.g., 'file (1).txt') to the new file.</li>
                    <li><b>Skip:</b> The file causing the conflict is not moved, copied, or renamed.</li></ul>
            </li>
            <li><b>Scan Subfolders Checkbox:</b> If checked, the 'Organize' preview will include files from all subdirectories within the selected folder (except those inside category folders like 'Images', 'Documents', etc.).</li>
            <li><b>Backup First Checkbox:</b> If checked, creates a `.zip` archive of the selected folder's contents <u>before</u> executing an 'Organize' action. The backup is placed in the parent directory of the selected folder.</li>
            <li><b>Watch Toggle & Label:</b> (Requires 'watchdog' library) If toggled on, automatically applies rules/categories to any <u>new</u> file created directly within the selected folder. Does not watch subfolders. Most other controls are disabled while watching.</li>
        </ul>

        <b style='font-size: 11pt;'>Main Area:</b><br>
        <ul>
            <li><b>Stats Label:</b> Shows the total file count and size for the selected folder (after clicking Browse or dropping a folder).</li>
            <li><b>Export Log Button:</b> Saves the text currently visible in the Log Area to a `.txt` file.</li>
            <li><b>Log Area:</b> Displays messages about ongoing processes, results, warnings, and errors.</li>
            <li><b>Progress Bar:</b> Shows the progress of tasks like organizing, scanning, or cleaning.</li>
        </ul>

        <b style='font-size: 11pt;'>Window Controls & System Tray:</b><br>
        <ul>
            <li><b>Theme Toggle (Sun/Moon Icon):</b> Switches between Light and Dark themes.</li>
            <li><b>Help Button (? Icon):</b> Shows this manual.</li>
            <li><b>Minimize/Maximize/Close:</b> Standard window controls. Closing the window minimizes it to the system tray.</li>
            <li><b>System Tray Icon (Folder):</b> Right-click to 'Show' the window again or 'Quit' the application completely. Notifications about background tasks (like Watcher actions) appear here.</li>
        </ul>
        """
        msgBox = QMessageBox(self)
        msgBox.setIcon(QMessageBox.Icon.Information)
        msgBox.setWindowTitle("File Organizer Manual")
        msgBox.setTextFormat(Qt.TextFormat.RichText)
        msgBox.setText(manual_text)
        msgBox.setStandardButtons(QMessageBox.StandardButton.Ok)
        msgBox.exec()


class CategoryEditor(QDialog):
    """Dialog for editing file type categories."""

    def __init__(self, categories, parent=None):
        super().__init__(parent)
        self.categories = json.loads(json.dumps(categories))  # Deep copy
        self.setWindowTitle("Category Editor")
        self.setMinimumSize(600, 400)

        layout = QVBoxLayout(self)
        self.table = QTableWidget()
        self.table.setColumnCount(2)
        self.table.setHorizontalHeaderLabels(["Category Name", "Extensions (comma-separated, e.g. .jpg,.png)"])
        self.table.horizontalHeader().setSectionResizeMode(1, QHeaderView.ResizeMode.Stretch)
        self.table.setColumnWidth(0, 150)

        self.refresh_table()
        layout.addWidget(self.table)

        btn_layout = QHBoxLayout()
        add_btn = QPushButton(qta.icon("fa5s.plus"), " Add Category")
        add_btn.clicked.connect(self.add_category)
        remove_btn = QPushButton(qta.icon("fa5s.trash-alt"), " Remove Category")
        remove_btn.clicked.connect(self.remove_category)
        save_btn = QPushButton("Save & Close")
        save_btn.setObjectName("SaveButton")
        save_btn.clicked.connect(self.accept)

        btn_layout.addWidget(add_btn)
        btn_layout.addWidget(remove_btn)
        btn_layout.addStretch()
        btn_layout.addWidget(save_btn)
        layout.addLayout(btn_layout)

    def refresh_table(self):
        self.table.setRowCount(len(self.categories))
        for i, (name, exts) in enumerate(self.categories.items()):
            name_item = QTableWidgetItem(name)
            if name == "Other":  # Don't allow renaming "Other" category
                name_item.setFlags(name_item.flags() & ~Qt.ItemFlag.ItemIsEditable)

            ext_str = ", ".join(exts)
            ext_item = QTableWidgetItem(ext_str)

            self.table.setItem(i, 0, name_item)
            self.table.setItem(i, 1, ext_item)

    def save_categories_from_table(self):
        """Reads the table and updates the self.categories dict."""
        new_categories = {}
        for row in range(self.table.rowCount()):
            name_item = self.table.item(row, 0)
            ext_item = self.table.item(row, 1)
            if not name_item or not ext_item:
                continue

            name = name_item.text().strip()
            if not name:
                QMessageBox.warning(self, "Invalid Name", f"Category name in row {row+1} cannot be empty.")
                return False  # Indicate failure

            ext_str = ext_item.text().strip()
            # Normalize extensions: split by comma, strip spaces, add dot if missing, remove empty
            ext_list = [f".{e.strip().lstrip('.')}" for e in ext_str.split(",") if e.strip()]

            if name in new_categories:
                QMessageBox.warning(self, "Duplicate Name", f"Category name '{name}' is duplicated.")
                return False  # Indicate failure

            new_categories[name] = sorted(list(set(ext_list)))  # Remove duplicate extensions

        if "Other" not in new_categories:
            new_categories["Other"] = []  # Ensure "Other" category always exists

        self.categories = new_categories
        return True  # Indicate success

    def add_category(self):
        self.save_categories_from_table()  # Save current edits first
        new_name = "New Category"
        count = 1
        while new_name in self.categories:
            new_name = f"New Category ({count})"
            count += 1
        self.categories[new_name] = []
        self.refresh_table()
        self.table.scrollToBottom()

    def remove_category(self):
        row = self.table.currentRow()
        if row < 0:
            return

        name = self.table.item(row, 0).text()
        if name == "Other":
            QMessageBox.warning(self, "Cannot Remove", "The 'Other' category cannot be removed.")
            return

        reply = QMessageBox.question(self, "Confirm Delete", f"Delete category '{name}'?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            self.save_categories_from_table()
            if name in self.categories:
                del self.categories[name]
            self.refresh_table()

    def accept(self):
        """Override accept to save before closing."""
        if self.save_categories_from_table():
            super().accept()

    def get_categories(self):
        return self.categories


class DuplicateFilesDialog(QDialog):
    def __init__(self, duplicates, parent=None):
        super().__init__(parent)
        self.duplicates = duplicates
        self.setWindowTitle("Duplicate File Finder")
        self.setMinimumSize(800, 600)
        self.deleted_count = 0
        layout = QVBoxLayout(self)
        self.table = QTableWidget()
        self.table.setColumnCount(3)
        self.table.setHorizontalHeaderLabels(["", "File Name", "Path"])
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.setColumnWidth(0, 40)
        self.populate_table()
        layout.addWidget(self.table)
        btn_layout = QHBoxLayout()
        self.status_label = QLabel("Select files to delete.")
        delete_btn = QPushButton(qta.icon("fa5s.trash-alt", color="white"), " Delete Selected")
        delete_btn.setObjectName("DeleteButton")
        delete_btn.clicked.connect(self.delete_selected)
        btn_layout.addWidget(self.status_label)
        btn_layout.addStretch()
        btn_layout.addWidget(delete_btn)
        layout.addLayout(btn_layout)

    def populate_table(self):
        self.table.setRowCount(sum(len(files) for files in self.duplicates.values()))
        row = 0
        for hash_val, files in self.duplicates.items():
            for i, filepath in enumerate(files):
                checkbox = QCheckBox()
                if i == 0:
                    checkbox.setDisabled(True)
                else:
                    checkbox.setChecked(True)
                self.table.setCellWidget(row, 0, checkbox)
                self.table.setItem(row, 1, QTableWidgetItem(os.path.basename(filepath)))
                self.table.setItem(row, 2, QTableWidgetItem(filepath))
                row += 1

    def delete_selected(self):
        files_to_delete = [
            self.table.item(row, 2).text()
            for row in range(self.table.rowCount())
            if self.table.cellWidget(row, 0) and self.table.cellWidget(row, 0).isChecked()
        ]
        if not files_to_delete:
            QMessageBox.warning(self, "No files selected", "Please select files to delete.")
            return
        reply = QMessageBox.question(self, "Confirm Deletion", f"Are you sure you want to permanently delete {len(files_to_delete)} files?", QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if reply == QMessageBox.StandardButton.Yes:
            error_count = 0
            self.deleted_count = 0
            for filepath in files_to_delete:
                try:
                    os.remove(filepath)
                    self.deleted_count += 1
                except OSError:
                    self.status_label.setText(f"Error deleting {os.path.basename(filepath)}.")
                    error_count += 1
            if error_count > 0:
                QMessageBox.warning(self, "Deletion Errors", f"Deleted {self.deleted_count} files, failed to delete {error_count} (check permissions).")
            else:
                QMessageBox.information(self, "Success", f"Successfully deleted {self.deleted_count} files.")
            self.accept()


class PreviewDialog(QDialog):
    """Dialog to show proposed file moves and get confirmation."""

    def __init__(self, proposed_actions, parent=None):  # Changed argument name
        super().__init__(parent)
        self.setWindowTitle("Organization Preview")
        self.setMinimumSize(800, 500)

        # Get base path from parent window
        base_path = ""
        if parent and hasattr(parent, "path_input"):
            base_path = parent.path_input.text()

        layout = QVBoxLayout(self)
        layout.addWidget(QLabel(f"<b>Proposed Changes ({len(proposed_actions)} files):</b>"))

        self.table = QTableWidget()
        self.table.setColumnCount(3)  # Action Type column
        self.table.setHorizontalHeaderLabels(["Original File", "Action", "Details / New Location"])
        self.table.horizontalHeader().setSectionResizeMode(2, QHeaderView.ResizeMode.Stretch)
        self.table.setColumnWidth(0, 250)
        self.table.setColumnWidth(1, 80)
        self.table.setRowCount(len(proposed_actions))
        self.table.setAlternatingRowColors(True)
        self.table.setSortingEnabled(True)  # Make sortable

        delete_font = QFont()
        delete_font.setStrikeOut(True)
        delete_color = QColor("#f85149" if parent.current_theme == "dark" else "#cf222e")
        copy_color = QColor("#3fb950" if parent.current_theme == "dark" else "#1f883d")

        for i, (action_type, original_full_path, details) in enumerate(proposed_actions):
            # Show relative path if possible, otherwise just basename
            display_name = os.path.basename(original_full_path)
            if base_path:
                try:
                    display_name = os.path.relpath(original_full_path, base_path)
                except ValueError:
                    pass  # Keep basename if relpath fails (e.g., different drive)

            item_original = QTableWidgetItem(display_name)
            item_action = QTableWidgetItem(action_type.capitalize())
            item_details = QTableWidgetItem(details)

            if action_type == "delete":
                item_original.setFont(delete_font)
                item_original.setForeground(delete_color)
                item_action.setForeground(delete_color)
                item_details.setForeground(delete_color)
            elif action_type == "copy":
                item_action.setForeground(copy_color)

            self.table.setItem(i, 0, item_original)
            self.table.setItem(i, 1, item_action)
            self.table.setItem(i, 2, item_details)

        layout.addWidget(self.table)

        button_box = QHBoxLayout()
        approve_btn = QPushButton(qta.icon("fa5s.check", color="white"), " Approve & Organize")
        approve_btn.setObjectName("ApproveButton")
        approve_btn.clicked.connect(self.accept)
        cancel_btn = QPushButton("Cancel")
        cancel_btn.clicked.connect(self.reject)
        button_box.addStretch()
        button_box.addWidget(cancel_btn)
        button_box.addWidget(approve_btn)
        layout.addLayout(button_box)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    if not QSystemTrayIcon.isSystemTrayAvailable():
        QMessageBox.critical(None, "Error", "No system tray detected on this system.")
        sys.exit(1)

    window = FileOrganizerWindow()
    window.show()
    sys.exit(app.exec())
